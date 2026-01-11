from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.image.exceptions import UnrecognizedImageError

from tg_pack.copys import copy_file_name
from tg_pack.find_links_service import find_links_excluding_platforms, extract_vk_video_links, find_rutube_links, \
    find_youtube_links
from tg_pack.find_messages import day_message_find, from_name_find, reply_to_find, text_find, media_wrap_find, \
    process_text_with_links
from tg_pack.gets_previews import get_other_preview, get_rt_preview, get_yt_preview
from tg_pack.hyperlinks import add_file_hyperlink

GR = '\033[92m'
RD = '\033[91m'
DYEL = '\033[33m'
BLD = '\033[1m'
UNDLN = '\033[4m'
DFLT = '\033[0m'
BL = '\033[94m'
STAND = '\033[39m'


def dict_worker(message, soup):
    mess_dict = {
        "day_message": None,
        "from_name": None,
        "reply_to": None,
        "text": None,
        "media_wrap": None
    }
    if day_message := day_message_find(message):
        mess_dict["day_message"] = day_message
    if from_name := from_name_find(message):
        mess_dict["from_name"] = from_name
    if reply_to := reply_to_find(message, soup):
        mess_dict["reply_to"] = reply_to
    if text := text_find(message):
        mess_dict["text"] = text
    if media_wrap := media_wrap_find(message):
        mess_dict["media_wrap"] = media_wrap
    return mess_dict


def mess_dict_work(mess_dict, doc, path):
    for key, val in mess_dict.items():
        if key == "from_name":
            if val:
                val_name_date = f"{val} | {mess_dict['day_message']}"
                doc.add_heading(val_name_date, level=1)
        elif key == "reply_to":
            if val:
                doc.add_paragraph(str(val), style='Intense Quote')
        elif key == "text":
            links_all = []
            if val:
                if found_yt_links := find_youtube_links(str(val)):
                    for yt_link in found_yt_links:
                        if yt_link not in links_all:
                            print(f"\n{GR}Получение превью сайта:{STAND}\n{yt_link}")
                            links_all.append(yt_link)
                            img_path = get_yt_preview(path, str(yt_link))
                            if img_path:
                                try:
                                    doc.add_picture(img_path, width=Inches(3))
                                    Path(img_path).unlink()
                                except UnrecognizedImageError:
                                    Path(img_path).unlink()
                if found_rt_links := find_rutube_links(str(val)):
                    for rt_link in found_rt_links:
                        if rt_link not in links_all:
                            print(f"\n{GR}Получение превью сайта:{STAND}\n{rt_link.strip(":").strip()}")
                            links_all.append(rt_link.strip(":").strip())
                            img_path = get_rt_preview(path, str(rt_link.strip(":").strip()))
                            if img_path:
                                try:
                                    doc.add_picture(img_path, width=Inches(3))
                                    Path(img_path).unlink()
                                except UnrecognizedImageError:
                                    Path(img_path).unlink()
                if found_vkvideo_links := extract_vk_video_links(str(val)):
                    for vkvideo_link in found_vkvideo_links:
                        if vkvideo_link not in links_all:
                            print(f"\n{GR}Получение превью сайта:{STAND}\n{vkvideo_link.strip(":").strip()}")
                            links_all.append(vkvideo_link.strip(":").strip())
                            img_path, art_title = get_other_preview(path, str(vkvideo_link.strip(":").strip()))
                            if img_path:
                                try:
                                    doc.add_picture(img_path, width=Inches(3))
                                    Path(img_path).unlink()
                                except UnrecognizedImageError:
                                    Path(img_path).unlink()
                            if art_title:
                                doc.add_paragraph(art_title, style='Intense Quote')
                if links_platforms := find_links_excluding_platforms(str(val)):
                    for link in links_platforms:
                        if link not in links_all:
                            print(f"\n{GR}Получение превью сайта:{STAND}\n{link}")
                            links_all.append(link)
                            img_path, art_title = get_other_preview(path, str(link.strip(":").strip()))
                            if img_path:
                                try:
                                    doc.add_picture(img_path, width=Inches(3))
                                    Path(img_path).unlink()
                                except UnrecognizedImageError:
                                    Path(img_path).unlink()
                            if art_title:
                                doc.add_paragraph(art_title, style='Intense Quote')
                process_text_with_links(doc, str(val))
        elif key == "media_wrap":
            if val:
                try:
                    doc.add_picture(str(Path(path) / str(val)), width=Inches(5))
                except (UnrecognizedImageError, FileNotFoundError):
                    if str(val).lower().endswith(".tgs"):
                        doc.add_paragraph("Stickers", style='Intense Quote')
                        continue
                    if str(val).lower().endswith((".mp4", ".mov", ".webm")):
                        name_v = f'{Path(str(val)).name}_thumb.jpg'
                        if (Path(path) / Path(str(val)).parent / name_v).exists():
                            doc.add_picture(str(Path(path) / Path(str(val)).parent / name_v), width=Inches(2))
                            folders = "video"
                        else:
                            folders = "files"
                        if rel_path := copy_file_name(Path(str(val)).name, path, val, folders):
                            p1 = doc.add_paragraph()
                            name_f = str(Path(rel_path).name)
                            add_file_hyperlink(p1, f"Открыть файл: {name_f}", rel_path, is_relative=True)


def parse(soup, path, name):
    doc = Document()
    print(f"\n{GR}Processing: {STAND}{name}.html\n{GR}{'*'*30}{STAND}")

    message_default_clearfix_lst = []
    if message_default_clearfix := soup.find_all("div", {"class": "message default clearfix"}):
        message_default_clearfix_lst.extend(message_default_clearfix)
    if message_default_clearfix_joined := soup.find_all("div", {"class": "message default clearfix joined"}):
        message_default_clearfix_lst.extend(message_default_clearfix_joined)

    sorted_messages = sorted(message_default_clearfix_lst, key=lambda x: int(x['id'].replace('message', '')))
    len_messages = len(sorted_messages)

    for nm, message in enumerate(sorted_messages, 1):
        mess_dict = dict_worker(message, soup)
        mess_dict_work(mess_dict, doc, path)
        print(f"\r{DYEL}[{STAND}{nm}{GR}/{STAND}{len_messages}{DYEL}] {STAND}| {GR}Message processed{STAND}", end="")

    doc.save(str(Path(path) / f"{name}.docx"))
    print(f"{GR}\n{'*' * 30}\n{DYEL}Document saving{STAND}")
